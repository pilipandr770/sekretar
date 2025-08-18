"""Document processing service for extracting text from various file formats."""
import os
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from urllib.parse import urlparse
import requests
from io import BytesIO
import logging

# Document processing libraries (with graceful fallbacks)
try:
    import PyPDF2
    HAS_PDF_SUPPORT = True
except ImportError:
    PyPDF2 = None
    HAS_PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX_SUPPORT = True
except ImportError:
    DocxDocument = None
    HAS_DOCX_SUPPORT = False

try:
    from bs4 import BeautifulSoup
    HAS_HTML_SUPPORT = True
except ImportError:
    BeautifulSoup = None
    HAS_HTML_SUPPORT = False

try:
    import tiktoken
    HAS_TIKTOKEN = True
except ImportError:
    tiktoken = None
    HAS_TIKTOKEN = False

from flask import current_app
from app.utils.exceptions import ProcessingError


logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing documents and extracting text content."""
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'md': 'text/markdown',
        'txt': 'text/plain',
        'html': 'text/html',
        'htm': 'text/html'
    }
    
    # Maximum file size (16MB)
    MAX_FILE_SIZE = 16 * 1024 * 1024
    
    # Text chunking parameters
    DEFAULT_CHUNK_SIZE = 1000  # tokens
    DEFAULT_OVERLAP = 200      # tokens
    
    @classmethod
    def extract_text_from_file(cls, file_path: str, mime_type: str = None) -> Dict[str, Any]:
        """
        Extract text content from a file.
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file (optional, will be detected)
            
        Returns:
            Dict containing extracted text, metadata, and processing info
        """
        try:
            # Validate file exists and size
            if not os.path.exists(file_path):
                raise ProcessingError(f"File not found: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > cls.MAX_FILE_SIZE:
                raise ProcessingError(f"File too large: {file_size} bytes (max: {cls.MAX_FILE_SIZE})")
            
            # Detect MIME type if not provided
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(file_path)
            
            # Get file extension
            file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            
            # Extract text based on file type
            if file_ext == 'pdf' or mime_type == 'application/pdf':
                content, metadata = cls._extract_pdf_text(file_path)
            elif file_ext in ['doc', 'docx'] or 'word' in (mime_type or ''):
                content, metadata = cls._extract_docx_text(file_path)
            elif file_ext == 'md' or mime_type == 'text/markdown':
                content, metadata = cls._extract_markdown_text(file_path)
            elif file_ext in ['html', 'htm'] or mime_type == 'text/html':
                content, metadata = cls._extract_html_text(file_path)
            elif file_ext == 'txt' or mime_type == 'text/plain':
                content, metadata = cls._extract_plain_text(file_path)
            else:
                # Try to read as plain text
                try:
                    content, metadata = cls._extract_plain_text(file_path)
                except UnicodeDecodeError:
                    raise ProcessingError(f"Unsupported file type: {file_ext} ({mime_type})")
            
            # Calculate token count
            token_count = cls._count_tokens(content)
            
            # Generate content hash
            content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
            
            return {
                'content': content,
                'token_count': token_count,
                'content_hash': content_hash,
                'file_size': file_size,
                'mime_type': mime_type,
                'file_extension': file_ext,
                'metadata': metadata
            }
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to extract text: {str(e)}")
    
    @classmethod
    def extract_text_from_url(cls, url: str, max_depth: int = 1) -> List[Dict[str, Any]]:
        """
        Extract text content from a URL.
        
        Args:
            url: URL to crawl
            max_depth: Maximum crawling depth (not implemented yet)
            
        Returns:
            List of extracted content dictionaries
        """
        try:
            # Validate URL
            if not cls._is_valid_url(url):
                raise ProcessingError(f"Invalid URL: {url}")
            
            # Make HTTP request
            headers = {
                'User-Agent': 'AI-Secretary-Bot/1.0 (Document Processor)',
                'Accept': 'text/html,application/xhtml+xml,text/plain,application/pdf'
            }
            
            response = requests.get(url, headers=headers, timeout=30, stream=True)
            response.raise_for_status()
            
            # Check content length
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > cls.MAX_FILE_SIZE:
                raise ProcessingError(f"Content too large: {content_length} bytes")
            
            # Get content type
            content_type = response.headers.get('content-type', '').lower()
            
            # Read content with size limit
            content_bytes = b''
            for chunk in response.iter_content(chunk_size=8192):
                content_bytes += chunk
                if len(content_bytes) > cls.MAX_FILE_SIZE:
                    raise ProcessingError("Content too large")
            
            # Process based on content type
            if 'text/html' in content_type:
                content, metadata = cls._extract_html_content(content_bytes, url)
            elif 'text/plain' in content_type:
                content = content_bytes.decode('utf-8', errors='ignore')
                metadata = {'content_type': content_type}
            elif 'application/pdf' in content_type:
                content, metadata = cls._extract_pdf_content(content_bytes)
            else:
                # Try to decode as text
                try:
                    content = content_bytes.decode('utf-8', errors='ignore')
                    metadata = {'content_type': content_type}
                except Exception:
                    raise ProcessingError(f"Unsupported content type: {content_type}")
            
            # Calculate token count
            token_count = cls._count_tokens(content)
            
            # Generate content hash
            content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
            
            # Extract title from URL or content
            title = metadata.get('title', cls._extract_title_from_url(url))
            
            return [{
                'content': content,
                'title': title,
                'url': url,
                'token_count': token_count,
                'content_hash': content_hash,
                'content_type': content_type,
                'metadata': metadata
            }]
            
        except requests.RequestException as e:
            logger.error(f"Failed to fetch URL {url}: {str(e)}")
            raise ProcessingError(f"Failed to fetch URL: {str(e)}")
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"Failed to process URL {url}: {str(e)}")
            raise ProcessingError(f"Failed to process URL: {str(e)}")
    
    @classmethod
    def chunk_text(cls, text: str, chunk_size: int = None, overlap: int = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks with overlap for context preservation.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum tokens per chunk
            overlap: Number of tokens to overlap between chunks
            
        Returns:
            List of chunk dictionaries
        """
        try:
            # Use TextChunker for better chunking logic
            from app.services.text_chunker import TextChunker, ChunkConfig
            
            config = ChunkConfig(
                chunk_size=chunk_size or cls.DEFAULT_CHUNK_SIZE,
                overlap=overlap or cls.DEFAULT_OVERLAP
            )
            
            chunker = TextChunker(config)
            return chunker.chunk_text(text, config)
            
        except Exception as e:
            logger.error(f"Failed to chunk text: {str(e)}")
            raise ProcessingError(f"Failed to chunk text: {str(e)}")
    
    @staticmethod
    def _extract_pdf_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        if not HAS_PDF_SUPPORT:
            raise ProcessingError("PDF support not available. Please install PyPDF2.")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else ''
                }
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
                
                content = '\n\n'.join(text_content)
                
                if not content.strip():
                    raise ProcessingError("No text content found in PDF")
                
                return content, metadata
                
        except Exception as e:
            raise ProcessingError(f"Failed to extract PDF text: {str(e)}")
    
    @staticmethod
    def _extract_pdf_content(content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF content bytes."""
        if not HAS_PDF_SUPPORT:
            raise ProcessingError("PDF support not available. Please install PyPDF2.")
        
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(content_bytes))
            
            # Extract metadata
            metadata = {
                'pages': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
            }
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception:
                    continue
            
            content = '\n\n'.join(text_content)
            
            if not content.strip():
                raise ProcessingError("No text content found in PDF")
            
            return content, metadata
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract PDF content: {str(e)}")
    
    @staticmethod
    def _extract_docx_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        if not HAS_DOCX_SUPPORT:
            raise ProcessingError("DOCX support not available. Please install python-docx.")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            content = '\n\n'.join(paragraphs)
            
            if not content.strip():
                raise ProcessingError("No text content found in document")
            
            return content, metadata
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract DOCX text: {str(e)}")
    
    @staticmethod
    def _extract_markdown_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic metadata extraction
            lines = content.split('\n')
            title = ''
            
            # Look for title in first few lines
            for line in lines[:10]:
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
            
            metadata = {
                'title': title,
                'lines': len(lines),
                'format': 'markdown'
            }
            
            return content, metadata
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract Markdown text: {str(e)}")
    
    @staticmethod
    def _extract_html_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            return DocumentProcessor._parse_html_content(html_content)
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract HTML text: {str(e)}")
    
    @staticmethod
    def _extract_html_content(content_bytes: bytes, url: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML content bytes."""
        try:
            # Decode content
            html_content = content_bytes.decode('utf-8', errors='ignore')
            
            content, metadata = DocumentProcessor._parse_html_content(html_content)
            metadata['url'] = url
            
            return content, metadata
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract HTML content: {str(e)}")
    
    @staticmethod
    def _parse_html_content(html_content: str) -> Tuple[str, Dict[str, Any]]:
        """Parse HTML content and extract text."""
        if not HAS_HTML_SUPPORT:
            raise ProcessingError("HTML support not available. Please install beautifulsoup4.")
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.get_text().strip() if title_tag else ''
        
        # Extract meta description
        description = ''
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            description = meta_desc.get('content', '').strip()
        
        # Extract main content
        # Try to find main content areas
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
        
        if main_content:
            text = main_content.get_text()
        else:
            # Fall back to body content
            body = soup.find('body')
            text = body.get_text() if body else soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        content = '\n'.join(line for line in lines if line)
        
        metadata = {
            'title': title,
            'description': description,
            'format': 'html'
        }
        
        return content, metadata
    
    @staticmethod
    def _extract_plain_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            metadata = {
                'lines': len(lines),
                'format': 'plain_text'
            }
            
            return content, metadata
            
        except Exception as e:
            raise ProcessingError(f"Failed to extract plain text: {str(e)}")
    
    @staticmethod
    def _count_tokens(text: str) -> int:
        """Count tokens in text using tiktoken."""
        if HAS_TIKTOKEN:
            try:
                encoding = tiktoken.get_encoding("cl100k_base")
                return len(encoding.encode(text))
            except Exception:
                pass
        
        # Fallback to approximate word count * 1.3
        return int(len(text.split()) * 1.3)
    
    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Validate URL format."""
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc])
        except Exception:
            return False
    
    @staticmethod
    def _extract_title_from_url(url: str) -> str:
        """Extract a title from URL."""
        try:
            parsed = urlparse(url)
            path = parsed.path.strip('/')
            
            if path:
                # Use last part of path as title
                title = path.split('/')[-1]
                # Remove file extension
                title = os.path.splitext(title)[0]
                # Replace hyphens and underscores with spaces
                title = title.replace('-', ' ').replace('_', ' ')
                # Capitalize words
                title = ' '.join(word.capitalize() for word in title.split())
                return title
            else:
                return parsed.netloc
                
        except Exception:
            return url